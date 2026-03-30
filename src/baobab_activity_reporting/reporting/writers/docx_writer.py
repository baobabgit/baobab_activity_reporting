"""Writer exportant un ReportModel au format Office Open XML (DOCX)."""

import logging
from pathlib import Path
from typing import Any, Mapping, cast

from docx import Document

from baobab_activity_reporting.exceptions.writer_error import WriterError
from baobab_activity_reporting.reporting.report_model import ReportModel
from baobab_activity_reporting.reporting.writers.abstract_writer import (
    AbstractWriter,
)

logger = logging.getLogger(__name__)


class DocxWriter(AbstractWriter):
    """Produit un fichier ``.docx`` à partir d'un :class:`ReportModel`.

    Utilise des styles Word standard : titre document, titres de section,
    paragraphes courants, légendes et tableaux. Aucune logique métier n'est
    appliquée au-delà du balisage structurel.

    :Example:
        >>> DocxWriter().output_format
        'docx'
    """

    @property
    def output_format(self) -> str:
        """Format logique ``docx``.

        :return: Identifiant de format.
        :rtype: str
        """
        return "docx"

    def write(self, model: ReportModel, output_path: Path | str) -> None:
        """Écrit le modèle au chemin indiqué.

        :param model: Rapport à exporter.
        :type model: ReportModel
        :param output_path: Fichier ``.docx`` cible.
        :type output_path: Path | str
        :raises WriterError: Si la création ou l'enregistrement échoue.
        """
        path = Path(output_path)
        try:
            doc = Document()
            self._render_document(doc, model)
            path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(path))
        except OSError as exc:
            raise WriterError(
                "Impossible d'écrire le fichier DOCX",
                output_format=self.output_format,
                details=str(exc),
            ) from exc
        except Exception as exc:  # pylint: disable=broad-exception-caught
            if isinstance(exc, WriterError):
                raise
            raise WriterError(
                "Échec du rendu DOCX",
                output_format=self.output_format,
                details=str(exc),
            ) from exc
        logger.info("DOCX enregistré : %s", path)

    def _render_document(self, doc: Any, model: ReportModel) -> None:
        """Remplit le document Word à partir du modèle.

        :param doc: Document python-docx vide ou partiel.
        :type doc: Any
        :param model: Source métier neutralisée.
        :type model: ReportModel
        """
        doc.add_heading(model.title, level=0)
        period_para = doc.add_paragraph()
        period_run = period_para.add_run(
            f"Période : {model.period_start} — {model.period_end}",
        )
        period_run.italic = True
        for block in model.preamble_narratives:
            self._add_body_paragraph(doc, block)
        for section in model.sections:
            self._render_section(doc, cast(Mapping[str, Any], section))

    def _render_section(
        self,
        doc: Any,
        section: Mapping[str, Any],
    ) -> None:
        """Ajoute une section (titre, narratifs, tableaux, insights).

        :param doc: Document en construction.
        :type doc: Any
        :param section: Bloc ``ReportBuilder`` tel quel.
        :type section: Mapping[str, Any]
        """
        title = str(section.get("title", section.get("section_code", "Section")))
        doc.add_heading(title, level=1)
        narratives = section.get("narrative_blocks")
        if isinstance(narratives, list):
            for para in narratives:
                self._add_body_paragraph(doc, str(para))
        tables = section.get("tables")
        if isinstance(tables, list):
            for tbl in tables:
                if isinstance(tbl, dict):
                    self._add_table(doc, cast(dict[str, Any], tbl))
        insights = section.get("insights")
        if isinstance(insights, list) and len(insights) > 0:
            doc.add_paragraph("Points saillants", style="Heading 3")
            for item in insights:
                doc.add_paragraph(str(item), style="List Bullet")

    def _add_table(  # pylint: disable=too-many-locals
        self,
        doc: Any,
        table_map: dict[str, Any],
    ) -> None:
        """Ajoute un tableau à partir de ``headers`` et ``rows``.

        :param doc: Document cible.
        :type doc: Any
        :param table_map: Structure ``TableBuilder``.
        :type table_map: dict[str, Any]
        """
        caption = table_map.get("caption")
        if caption is not None:
            cap = doc.add_paragraph()
            cap.add_run(str(caption)).bold = True
        headers = table_map.get("headers")
        rows = table_map.get("rows")
        if not isinstance(headers, list) or not isinstance(rows, list):
            return
        if len(headers) == 0:
            return
        n_rows = 1 + len(rows)
        n_cols = len(headers)
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            hdr_cells[idx].text = str(header)
        for r_idx, row in enumerate(rows):
            if not isinstance(row, (list, tuple)):
                continue
            row_seq: list[Any] = list(row)
            cells = table.rows[r_idx + 1].cells
            for c_idx, _header in enumerate(headers):
                text = row_seq[c_idx] if c_idx < len(row_seq) else ""
                cells[c_idx].text = str(text)

    @staticmethod
    def _add_body_paragraph(doc: Any, text: str) -> None:
        """Ajoute un paragraphe corps de texte.

        :param doc: Document cible.
        :type doc: Any
        :param text: Contenu.
        :type text: str
        """
        doc.add_paragraph(text)
